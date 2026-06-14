"""Konwersja Markdown → DOCX (python-docx)."""
from __future__ import annotations

import io
import re
from typing import TYPE_CHECKING

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
_LIST_RE = re.compile(r"^(\s*)([-*+]|\d+\.)\s+(.+)$")
_HR_RE = re.compile(r"^(-{3,}|\*{3,}|_{3,})\s*$")
_BLOCKQUOTE_RE = re.compile(r"^>\s?(.*)$")
_INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|_[^_]+_)")


def _add_inline_runs(paragraph: Paragraph, text: str) -> None:
    last = 0
    for match in _INLINE_RE.finditer(text):
        if match.start() > last:
            paragraph.add_run(text[last : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        last = match.end()
    if last < len(text):
        paragraph.add_run(text[last:])


def _configure_document_styles(document: docx.Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)


def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    """Zwraca plik .docx jako bajty."""
    document = docx.Document()
    _configure_document_styles(document)

    for raw_line in markdown_text.replace("\r\n", "\n").split("\n"):
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            document.add_paragraph("")
            continue

        if _HR_RE.match(stripped):
            document.add_paragraph("")
            continue

        heading = _HEADING_RE.match(stripped)
        if heading:
            level = min(len(heading.group(1)), 4)
            text = heading.group(2).strip()
            p = document.add_heading(text, level=level)
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        blockquote = _BLOCKQUOTE_RE.match(stripped)
        if blockquote:
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            _add_inline_runs(p, blockquote.group(1))
            continue

        list_match = _LIST_RE.match(line)
        if list_match:
            p = document.add_paragraph(style="List Bullet")
            _add_inline_runs(p, list_match.group(3).strip())
            continue

        p = document.add_paragraph()
        _add_inline_runs(p, stripped)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()

"""Template-based DOCX export for legal drafts using docxtpl."""
from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docxtpl import DocxTemplate, RichText


def _build_template_bytes() -> bytes:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15

    place_line = doc.add_paragraph("{{r place_date_rt}}")
    place_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph("")
    doc.add_paragraph("{{r recipient_rt}}")
    doc.add_paragraph("")

    title = doc.add_paragraph("{{r title_rt}}")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True

    doc.add_paragraph("")
    doc.add_paragraph("{{r body_rt}}")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("Z poważaniem,")
    doc.add_paragraph("{{r sender_rt}}")

    doc.add_paragraph("")
    doc.add_paragraph("Załączniki:")
    doc.add_paragraph("{{r attachments_rt}}")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _multiline_rich_text(value: str) -> RichText:
    rt = RichText()
    lines = [line.rstrip() for line in (value or "").replace("\r\n", "\n").split("\n")]
    if not lines:
        return rt

    for i, line in enumerate(lines):
        if i > 0:
            rt.add("\n")
        rt.add(line if line else " ")
    return rt


def render_draft_docx_bytes(
    *,
    title: str,
    body_markdown: str,
    structured_data: dict[str, Any] | None = None,
) -> bytes:
    data = structured_data or {}
    sender = str(data.get("sender") or "........................................")
    recipient = str(data.get("recipient") or "........................................")
    place_date = str(data.get("placeDate") or "...................., dnia ....................")
    attachments = str(data.get("attachments") or "1. ........................................")

    template_stream = io.BytesIO(_build_template_bytes())
    tpl = DocxTemplate(template_stream)
    tpl.render(
        {
            "title_rt": _multiline_rich_text(title or "PISMO"),
            "body_rt": _multiline_rich_text(body_markdown or ""),
            "sender_rt": _multiline_rich_text(sender),
            "recipient_rt": _multiline_rich_text(recipient),
            "place_date_rt": _multiline_rich_text(place_date),
            "attachments_rt": _multiline_rich_text(attachments),
        }
    )

    output = io.BytesIO()
    tpl.save(output)
    return output.getvalue()


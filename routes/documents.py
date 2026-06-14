import os
import asyncio
import time
from typing import Optional
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import Response
from models.request_models import (
    DocumentUploadResponse,
    DocumentAnalysisRequest,
    DraftRequest,
)
from services.document_service import index_document_to_supabase
from services.draft_document_catalog import get_document_type_hint
from utils.helpers import sanitize_filename
import io
import pypdf
import docx
from schemas.chat_contract import LegalSourceType

from fastapi import BackgroundTasks

router = APIRouter()

MAX_FILE_SIZE = 15 * 1024 * 1024  # 15MB — zgodne z limitem frontendu


@router.post("/upload-document", response_model=DocumentUploadResponse)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    category: str = Form("rag_user"),
    source_type: Optional[LegalSourceType] = Form(None),
):
    try:
        filename = sanitize_filename(file.filename or "unknown")
        print(f"\n   [UPLOAD] Otrzymano plik: {filename} ({file.content_type})")
        file_content = await file.read()
        print(f"   [UPLOAD] Rozmiar: {len(file_content)} bajtów. Rozpoczynanie ekstrakcji...")
        if len(file_content) > MAX_FILE_SIZE:
             print(f"   [UPLOAD ERROR] Plik zbyt duży ({len(file_content)} > {MAX_FILE_SIZE})")
             raise HTTPException(status_code=413, detail="Plik zbyt duży (maksymalnie 10MB)")
        
        os.makedirs("pdfs", exist_ok=True)
        with open(f"pdfs/{filename}", "wb") as f:
            f.write(file_content)
        
        print(f"   [UPLOAD] Plik zapisany lokalnie: {filename}. Przetwarzanie V2...")
        
        extracted_text = ""
        error = None
        pre_embedding = None

        try:
            # Lokalna ekstrakcja tekstu z kodu aplikacji
            if file.content_type == "application/pdf" or filename.lower().endswith(".pdf"):
                print(f"   [EXTRACT] Rozpoczynam odczyt PDF (bez skracania treści): {filename}")
                pdf_reader = pypdf.PdfReader(io.BytesIO(file_content))
                page_parts = []
                for i, page in enumerate(pdf_reader.pages, start=1):
                    page_text = page.extract_text()
                    if page_text is None:
                        page_text = ""
                    page_parts.append(f"--- STRONA {i} ---\n{page_text}")
                extracted_text = "\n\n".join(page_parts).strip()
                print(f"   [EXTRACT] Sukces PDF ({len(extracted_text)} znaków)")
                from services.ocr_cache import MIN_OCR_CHARS

                if len(extracted_text) < MIN_OCR_CHARS:
                    print(
                        f"   [EXTRACT WARN] PDF ma mało tekstu ({len(extracted_text)} zn.) — "
                        "możliwy skan; rozważ upload zdjęć stron (OCR wizyjny)."
                    )
                
            elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or filename.lower().endswith(".docx"):
                print(f"   [EXTRACT] Rozpoczynam odczyt DOCX do Markdown: {filename}")
                doc = docx.Document(io.BytesIO(file_content))
                md_lines = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    
                    style_name = para.style.name if para.style else ""
                    
                    if style_name.startswith('Heading'):
                        try:
                            level = int(style_name.replace('Heading ', ''))
                            md_lines.append(f"\n{'#' * level} {text}\n")
                        except:
                            md_lines.append(f"\n## {text}\n")
                    elif 'List' in style_name:
                        md_lines.append(f"- {text}")
                    else:
                        is_bold = all(run.bold for run in para.runs if run.text.strip())
                        if is_bold and text:
                            md_lines.append(f"\n**{text}**\n")
                        else:
                            md_lines.append(text)
                extracted_text = "\n".join(md_lines)
                print(f"   [EXTRACT] Sukces DOCX ({len(extracted_text)} znaków)")
                
            elif file.content_type == "text/plain" or filename.lower().endswith(".txt"):
                print(f"   [EXTRACT] Rozpoczynam odczyt TXT: {filename}")
                extracted_text = file_content.decode("utf-8", errors="ignore")
                
            elif (file.content_type and file.content_type.startswith("image/")) or filename.lower().endswith((".jpg", ".jpeg", ".png", ".webp")):
                print(f"   [EXTRACT] Rozpoczynam dosłowny OCR wizyjny obrazu: {filename}")
                try:
                    from moa.http_client import get_shared_openai_client
                    from services.ocr_cache import MIN_OCR_CHARS, set_cached_ocr_for_image
                    from services.vision_ocr import run_verbatim_vision_ocr

                    client = get_shared_openai_client()
                    extracted_text = ""
                    model_name = None
                    last_err = None
                    try:
                        extracted_text, model_name = await run_verbatim_vision_ocr(
                            client, file_content
                        )
                    except Exception as ocr_err:
                        last_err = ocr_err
                        extracted_text = ""

                    if extracted_text.strip() and len(extracted_text.strip()) < MIN_OCR_CHARS:
                        print(
                            f"   [EXTRACT] OCR za krótki ({len(extracted_text.strip())} < {MIN_OCR_CHARS} zn.)"
                        )
                        extracted_text = ""

                    if extracted_text.strip():
                        print(
                            f"   [EXTRACT] Sukces OCR (model={model_name}, {len(extracted_text)} znaków)"
                        )
                        set_cached_ocr_for_image(file_content, extracted_text)
                    elif last_err:
                        raise last_err
                    else:
                        raise RuntimeError("Wszystkie modele wizyjne zawiodły lub OCR zbyt krótki")
                except Exception as vision_err:
                    print(f"   [EXTRACT ERR] Błąd analizy wizyjnej obrazu: {vision_err}")
                    extracted_text = f"[Plik graficzny {filename}. Nie udało się wyekstrahować tekstu podczas przesyłania: {vision_err}]"
            else:
                extracted_text = f"[V2: Plik binarny {filename}]"
                
        except Exception as extract_err:
            error = f"Błąd ekstrakcji: {extract_err}"
            print(f"   [EXTRACT ERROR] {error}")
        
        success = not bool(error)
        
        # Automatycznie zapisz załącznik w bazie, by zachować wyniki OCR i umożliwić dostęp z Biblioteki
        if success and extracted_text:
            from services.document_service import index_document_to_supabase
            
            async def background_indexing():
                try:
                    await index_document_to_supabase(
                        file_content=file_content,
                        filename=filename,
                        content_type=file.content_type or "",
                        category=category,
                        pre_extracted_text=extracted_text,
                        pre_embedding=pre_embedding,
                        source_type=(source_type.value if source_type else None),
                    )
                    table_name = "knowledge_base_legal" if category == "rag_legal" else "knowledge_base_user"
                    print(f"   [BACKGROUND] Zapisano dokument {filename} w bazie ({table_name}).")
                except Exception as e:
                    print(f"   [BACKGROUND ERROR] Błąd podczas indeksowania {filename}: {e}")

            background_tasks.add_task(background_indexing)

        return DocumentUploadResponse(
            success=success, filename=filename, 
            extracted_text=extracted_text if success else "", 
            text_length=len(extracted_text) if success else 0,
            error=error
        )
    except Exception as e:
        return DocumentUploadResponse(success=False, filename="unknown", extracted_text="", text_length=0, error=str(e))

@router.post("/upload")
@router.post("/index-document")
async def index_document_to_rag(
    file: UploadFile = File(...), 
    category: str = Form("rag_legal"),
    source_type: Optional[LegalSourceType] = Form(None),
):
    file_content = await file.read()
    if len(file_content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="Plik zbyt duży (maksymalnie 10MB)")
    
    os.makedirs("pdfs", exist_ok=True)
    filename = sanitize_filename(file.filename or "unknown")
    with open(f"pdfs/{filename}", "wb") as f:
        f.write(file_content)

    return await index_document_to_supabase(
        file_content,
        filename,
        file.content_type or "",
        category=category,
        source_type=(source_type.value if source_type else None),
    )

@router.post("/index-saved-file/{filename}")
async def index_saved_file(filename: str):
    filename = sanitize_filename(filename)
    file_path = f"pdfs/{filename}"
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Plik nie istnieje na serwerze")
        
    with open(file_path, "rb") as f:
        file_content = f.read()

    return await index_document_to_supabase(file_content, filename, "")

@router.post("/export-docx")
async def export_docx(request: DocumentAnalysisRequest):
    """Eksport wygenerowanego pisma (Markdown) do pliku Word (.docx)."""
    if not request.document_text:
        raise HTTPException(status_code=400, detail="Brak treści pisma")

    from services.docx_export import markdown_to_docx_bytes
    from services.docx_template_export import render_draft_docx_bytes

    base_name = sanitize_filename(request.question or "pismo")
    if base_name.lower().endswith(".md"):
        base_name = base_name[:-3]
    if base_name.lower().endswith(".docx"):
        base_name = base_name[:-5]
    filename = f"{base_name}.docx"

    try:
        docx_bytes = render_draft_docx_bytes(
            title=request.question,
            body_markdown=request.document_text,
            structured_data=request.structured_data,
        )
    except Exception as template_err:
        print(f"   [DOCX TEMPLATE] fallback to markdown exporter: {template_err}")
        docx_bytes = markdown_to_docx_bytes(request.document_text)

    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _build_draft_user_prompt(request: DraftRequest, rag_context: str = "") -> str:
    parts: list[str] = []

    document_hint = get_document_type_hint(request.document_type)
    if document_hint:
        parts.append(f"### TYP PISMA:\n{document_hint}\n\n")

    structured = request.structured_data or {}
    if structured:
        parts.append("### DANE STRUKTURALNE DO DOKUMENTU:\n")
        parts.append(
            f"[MIEJSCE I DATA]: {structured.get('placeDate') or '...................., dnia ....................'}\n\n"
        )
        parts.append(f"[NADAWCA]:\n{structured.get('sender') or '....................'}\n\n")
        parts.append(f"[ADRESAT]:\n{structured.get('recipient') or '....................'}\n\n")
        parts.append("---\n\n")
    parts.append(
        f"### INSTRUKCJE DO TREŚCI:\n{request.user_instructions or 'Sporządź odpowiednie pismo procesowe/urzędowe na podstawie danych powyżej.'}\n\n"
    )
    if rag_context:
        parts.append(f"### KONTEKST Z BAZY WIEDZY RAG:\n{rag_context}\n\n")
    parts.append(
        "[WYMAGANIE]: Wygeneruj wyłącznie gotowy dokument w formacie Markdown. "
        "Zachowaj profesjonalny układ (miejsce na podpis, załączniki)."
    )
    return "".join(parts)


async def _fetch_draft_rag_context(query: str) -> str:
    try:
        import httpx
        from moa.config import SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY, OPENROUTER_API_KEY

        if not all([SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY, OPENROUTER_API_KEY]):
            return ""

        async with httpx.AsyncClient(timeout=60) as client:
            emb_res = await client.post(
                "https://openrouter.ai/api/v1/embeddings",
                headers={
                    "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": "openai/text-embedding-3-small",
                    "input": query or "draft document",
                },
            )
            if emb_res.status_code != 200:
                return ""
            emb_data = emb_res.json()
            embedding = (emb_data.get("data") or [{}])[0].get("embedding")
            if not embedding:
                return ""

            rpc_res = await client.post(
                f"{SUPABASE_URL}/rest/v1/rpc/match_knowledge",
                headers={
                    "Authorization": f"Bearer {SUPABASE_SERVICE_ROLE_KEY}",
                    "apikey": SUPABASE_SERVICE_ROLE_KEY,
                    "Content-Type": "application/json",
                },
                json={
                    "query_embedding": embedding,
                    "match_threshold": 0.35,
                    "match_count": 5,
                },
            )
            if rpc_res.status_code != 200:
                return ""
            matches = rpc_res.json()
            if not isinstance(matches, list) or not matches:
                return ""
            return "\n---\n".join(
                (m.get("content") or "") for m in matches if isinstance(m, dict)
            )
    except Exception as exc:
        print(f"   [DRAFT] RAG context skipped: {exc}")
        return ""


@router.post("/draft-document")
async def draft_document(request: DraftRequest):
    """Generator pism — lokalny fallback gdy Edge Function zwraca 401 lub jest niedostępna."""
    from moa.http_client import get_shared_openai_client

    rag_context = await _fetch_draft_rag_context(request.user_instructions)
    final_user_prompt = _build_draft_user_prompt(request, rag_context)

    history_messages = []
    for item in request.history or []:
        if not isinstance(item, dict):
            continue
        role = item.get("role")
        content = item.get("content")
        if role in ("user", "assistant", "system") and content:
            history_messages.append({"role": role, "content": str(content)})

    client = get_shared_openai_client()
    response = await client.chat.completions.create(
        model=request.model,
        messages=[
            {
                "role": "system",
                "content": request.system_prompt
                or "Jesteś ekspertem ds. pism prawnych i urzędowych w Polsce.",
            },
            *history_messages,
            {"role": "user", "content": final_user_prompt},
        ],
    )
    content = response.choices[0].message.content or ""
    if not content.strip():
        raise HTTPException(status_code=502, detail="Model nie zwrócił treści pisma.")
    return {"content": content}


@router.post("/save-draft")
async def save_draft(request: DocumentAnalysisRequest):
    """
    Zapisuje wygenerowane pismo do bazy użytkownika (RAG Ready).
    Wykorzystujemy DocumentAnalysisRequest bo ma pola 'name' (jako title) i 'content'.
    """
    if not request.document_text:
        raise HTTPException(status_code=400, detail="Brak treści pisma")
    
    filename = sanitize_filename(request.question or f"Pismo_{int(time.time())}.md")
    if not filename.endswith('.md'):
        filename += '.md'
        
    return await index_document_to_supabase(
        file_content=request.document_text.encode('utf-8'),
        filename=filename,
        content_type='text/markdown',
        category='rag_user'
    )

@router.get("/list")
async def list_documents():
    """
    Pobiera list wszystkich dokumentów z Supabase z metadanych
    """
    try:
        import httpx
        from moa.config import SUPABASE_URL, SUPABASE_ANON_KEY
        
        async with httpx.AsyncClient(timeout=30) as client:
            headers = {
                "Authorization": f"Bearer {SUPABASE_ANON_KEY}",
                "apikey": SUPABASE_ANON_KEY
            }
            
            # Pobierz z obu tabel (knowledge_base_legal i knowledge_base_user)
            res_legal = await client.get(
                f"{SUPABASE_URL}/rest/v1/knowledge_base_legal?select=metadata",
                headers=headers
            )
            res_user = await client.get(
                f"{SUPABASE_URL}/rest/v1/knowledge_base_user?select=metadata",
                headers=headers
            )
                
            documents = set()
            for response in [res_legal, res_user]:
                if response.status_code == 200:
                    data = response.json()
                    for item in data:
                        if isinstance(item, dict) and 'metadata' in item:
                            metadata = item['metadata']
                            if isinstance(metadata, dict) and 'filename' in metadata:
                                documents.add(metadata['filename'])
                
            return {
                "success": True,
                "documents": sorted(list(documents)),
                "count": len(documents)
            }
                
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "documents": []
        }

@router.get("/content/{filename}")
async def get_document_content(filename: str):
    """
    Pobiera zawartość dokumentu z lokalnego storage (obsługuje prefixy timestamp).
    """
    try:
        import glob
        safe_filename = sanitize_filename(filename)
        
        # Potencjalne lokalizacje i wzorce (w tym z prefixem timestamp)
        search_patterns = [
            f"local_storage/chat_attachments/*_{safe_filename}",
            f"local_storage/chat_attachments/{safe_filename}",
            f"local_storage/knowledge_base/*_{safe_filename}",
            f"local_storage/knowledge_base/{safe_filename}",
            f"local_storage/knowledge_base_legal/*_{safe_filename}",
            f"local_storage/knowledge_base_legal/{safe_filename}",
            f"pdfs/*_{safe_filename}",
            f"pdfs/{safe_filename}"
        ]
        
        found_path = None
        for pattern in search_patterns:
            matches = glob.glob(pattern)
            if matches:
                # Weź najnowszy (ostatni alfabetycznie przy timestampach)
                found_path = sorted(matches)[-1]
                break
        
        if found_path and os.path.exists(found_path):
            # Safe absolute path containment check to prevent Path Traversal
            abs_found = os.path.abspath(found_path)
            allowed_dirs = [
                os.path.abspath("local_storage"),
                os.path.abspath("pdfs")
            ]
            is_contained = False
            for d in allowed_dirs:
                if abs_found.startswith(d + os.sep) or abs_found == d:
                    is_contained = True
                    break
            if not is_contained:
                raise HTTPException(status_code=403, detail="Dostęp zabroniony (Wykryto Path Traversal)")

            with open(found_path, "rb") as f:
                file_content = f.read()
                print(f"   [CONTENT] Znaleziono plik: {found_path}")

                extracted_text = ""
                error = None
                path_lower = found_path.lower()
                is_raster_image = path_lower.endswith(
                    (".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp", ".heic", ".tiff", ".tif")
                )

                try:
                    if is_raster_image:
                        print(
                            "   [CONTENT] Obraz rastrowy — pomijam dekodowanie UTF-8; szukam tekstu OCR w Supabase (user → legal)."
                        )
                        error = None
                    elif path_lower.endswith(".pdf"):
                        pdf_reader = pypdf.PdfReader(io.BytesIO(file_content))
                        md_lines = []
                        for i, page in enumerate(pdf_reader.pages):
                            page_text = page.extract_text() or ""
                            for line in page_text.split('\n'):
                                line = line.strip()
                                if line:
                                    if len(line) < 80 and line.isupper():
                                        md_lines.append(f"\n## {line}\n")
                                    elif line.startswith(('•', '-', '*', '1.', '2.')):
                                        md_lines.append(f"{line}")
                                    else:
                                        md_lines.append(line)
                        extracted_text = "\n".join(md_lines)
                        error = None
                    elif found_path.lower().endswith(".docx"):
                        doc = docx.Document(io.BytesIO(file_content))
                        md_lines = []
                        for para in doc.paragraphs:
                            text = para.text.strip()
                            if text:
                                if para.style and para.style.name.startswith('Heading'):
                                    try:
                                        level = int(para.style.name.replace('Heading ', ''))
                                        md_lines.append(f"\n{'#' * level} {text}\n")
                                    except:
                                        md_lines.append(f"\n## {text}\n")
                                elif para.style and 'List' in para.style.name:
                                    md_lines.append(f"- {text}")
                                else:
                                    md_lines.append(text)
                        extracted_text = "\n".join(md_lines)
                        error = None
                    else:
                        extracted_text = file_content.decode("utf-8")
                        error = None
                except Exception as e:
                    extracted_text = ""
                    error = f"Błąd odczytu lokalnego pliku: {str(e)}"
                    print(f"   [CONTENT] {error}. Przejście do Supabase fallback...")
                
                pre_embedding = None
            
            if not error and extracted_text:
                return {
                    "success": True,
                    "filename": os.path.basename(found_path),
                    "content": extracted_text,
                    "size": len(file_content),
                    "path": found_path
                }
            # Remove the early return so it falls back to Supabase if local extraction fails
        
        # --- Supabase: tekst zindeksowany (OCR / PDF) — preferuj dokumenty użytkownika ---
        print(f"   [CONTENT] Szukanie treści w Supabase (filename={safe_filename})…")
        try:
            from moa.config import SUPABASE_URL, SUPABASE_ANON_KEY
            import httpx
            from urllib.parse import quote

            fn_eq = quote(safe_filename, safe="")
            
            async with httpx.AsyncClient(timeout=30) as client:
                headers = {
                    "Authorization": f"Bearer {SUPABASE_ANON_KEY}",
                    "apikey": SUPABASE_ANON_KEY
                }
                
                content_found = ""
                for table in ["knowledge_base_user", "knowledge_base_legal"]:
                    url = (
                        f"{SUPABASE_URL.rstrip('/')}/rest/v1/{table}"
                        f"?metadata->>filename=eq.{fn_eq}&select=content"
                    )
                    res = await client.get(url, headers=headers)
                    
                    if res.status_code == 200:
                        data = res.json()
                        if data:
                            # Concatenate all fragments
                            content_found = " ".join(
                                str(item.get("content") or "") for item in data
                            )
                            break
                
                if content_found:
                    return {
                        "success": True,
                        "filename": safe_filename,
                        "content": content_found,
                        "size": len(content_found),
                        "path": "supabase_record"
                    }
        except Exception as se:
            print(f"   [CONTENT ERROR] Supabase fallback failed: {se}")

        return {
            "success": False,
            "error": f"Dokument '{filename}' nie został znaleziony ani na dysku, ani w bazie danych",
            "filename": filename
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "filename": filename
        }

from fastapi import BackgroundTasks

@router.post("/index-knowledge-base")
async def trigger_full_indexing(background_tasks: BackgroundTasks):
    """
    Uruchamia pełne indeksowanie plików PDF z katalogu local_storage/knowledge_base w tle.
    """
    from moa.config import PROJECT_DIR
    folder = os.path.join(PROJECT_DIR, 'local_storage', 'knowledge_base')
    
    if not os.path.isdir(folder):
        return {"success": False, "error": f"Katalog nie istnieje: {folder}"}
        
    pdf_files = [f for f in os.listdir(folder) if f.lower().endswith('.pdf')]
    if not pdf_files:
        return {"success": False, "error": "Brak plików PDF do zindeksowania."}
        
    async def run_indexing():
        print(f"[START] [BG] Rozpoczynanie indeksowania {len(pdf_files)} plików...")
        for filename in sorted(pdf_files):
            try:
                path = os.path.join(folder, filename)
                if not os.path.exists(path): continue
                
                with open(path, 'rb') as f:
                    content = f.read()
                
                print(f"[INDEX] [BG] Indeksowanie: {filename}")
                result = await index_document_to_supabase(
                    file_content=content,
                    filename=filename,
                    content_type='application/pdf',
                    category='rag_legal'
                )
                
                if result.get('success'):
                    print(f"[OK] [BG] OK: {filename} ({result.get('fragments')} fragm.)")
                else:
                    print(f"[ERROR] [BG] Błąd {filename}: {result.get('error')}")
                    
            except Exception as e:
                print(f"[FATAL] [BG] Krytyczny błąd pliku {filename}: {e}")
        
        print("[FINISHED] [BG] Indeksowanie zakończone.")

    background_tasks.add_task(run_indexing)
    
    return {
        "success": True, 
        "message": f"Uruchomiono indeksowanie {len(pdf_files)} plików w tle. Sprawdź logi serwera API, aby śledzić postęp.",
        "files_count": len(pdf_files),
        "folder": folder
    }


@router.post("/analyze-document")
async def analyze_document_endpoint(request: DocumentAnalysisRequest):
    """
    Analizuje podany dokument w kontekście pytania użytkownika oraz (opcjonalnie) bazy wiedzy.
    """
    try:
        document_text = request.document_text or ""
        question = request.question or ""
        use_rag = request.use_rag
        
        # Ograniczenia i walidacja
        if not question.strip():
            return {
                "success": False,
                "answer": "Brak pytania do analizy.",
                "sources": [],
                "document_length": len(document_text),
                "context_length": 0,
                "rag_used": False,
                "error": "Pytanie nie może być puste."
            }

        rag_context = ""
        sources = []
        
        if use_rag:
            # 1. Retrieval z bazy ogólnej (knowledge_base_legal)
            try:
                from services.retrieval_service import retrieval_service
                legal_chunks = await retrieval_service.search_supabase(
                    query=question,
                    table_name="knowledge_base_legal",
                    match_count=5,
                    hybrid=True
                )
                for chunk in legal_chunks:
                    content = chunk.get("content") or ""
                    meta = chunk.get("metadata") or {}
                    filename = meta.get("filename") or "Baza prawna"
                    
                    rag_context += f"\n--- Źródło: {filename} ---\n{content}\n"
                    sources.append(filename)
            except Exception as e:
                print(f"[RAG LEGAL ERR] Błąd pobierania bazy legal: {e}")

        # Przygotowanie promptu do LLM
        system_prompt = (
            "Jesteś profesjonalnym polskim asystentem prawnym (LexMind). "
            "Twoim zadaniem jest merytoryczna, wyczerpująca i profesjonalna analiza dostarczonego dokumentu prawnego "
            "w kontekście pytania użytkownika oraz załączonej dodatkowej wiedzy prawnej."
        )

        user_content = f"""
ZADANIE: Analizuj poniższy dokument w kontekście pytania użytkownika.

PYTANIE UŻYTKOWNIKA:
{question}

TEKST DOKUMENTU:
{document_text}
"""

        if rag_context:
            user_content += f"""

DODATKOWA WIEDZA PRAWNA (RAG):
{rag_context}
"""

        user_content += "\n\nZwróć profesjonalną i merytoryczną odpowiedź w języku polskim."

        # Wywołanie LLM z naszym multi-model vision-like fallback loop
        from moa.config import get_async_client
        client = get_async_client()
        
        from config import settings
        models_to_try = list(settings.vision_ocr_models)
        
        answer = None
        error_msg = None
        
        for model in models_to_try:
            try:
                print(f"[ANALYZE] Próba generowania odpowiedzi modelem {model}...")
                completion = await client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_content}
                    ],
                    temperature=0.2
                )
                answer = completion.choices[0].message.content
                print(f"[ANALYZE OK] Odpowiedź wygenerowana pomyślnie przy użyciu modelu {model}.")
                break
            except Exception as e:
                error_msg = str(e)
                print(f"[ANALYZE ERR] Błąd modelu {model}: {e}")
                continue

        if not answer:
            raise Exception(f"Wszystkie modele LLM zwróciły błąd: {error_msg}")

        return {
            "success": True,
            "answer": answer,
            "sources": list(set(sources)),
            "document_length": len(document_text),
            "context_length": len(rag_context),
            "rag_used": len(rag_context) > 0
        }

    except Exception as e:
        print(f"[ANALYZE FATAL] Błąd krytyczny endpointu analyze-document: {e}")
        return {
            "success": False,
            "answer": "Wystąpił błąd krytyczny podczas analizy dokumentu.",
            "sources": [],
            "document_length": len(request.document_text) if request else 0,
            "context_length": 0,
            "rag_used": False,
            "error": str(e)
        }

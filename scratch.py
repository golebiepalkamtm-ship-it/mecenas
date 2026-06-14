import docx

doc = docx.Document(r"e:\moj prawnik\pdfs\Pismo_US_Luban_wezwanie_nadplaty.docx")
print("=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"[{i}] {p.text}")

print("=== TABLES ===")
for t_idx, table in enumerate(doc.tables):
    print(f"\nTable {t_idx}:")
    for r_idx, row in enumerate(table.rows):
        row_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        print(f"  Row {r_idx}: {row_text}")
